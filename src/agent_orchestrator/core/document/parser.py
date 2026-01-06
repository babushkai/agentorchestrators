"""Document parsing and text extraction."""

import io
from abc import ABC, abstractmethod
from typing import Any

import structlog
from pydantic import BaseModel, Field

logger = structlog.get_logger(__name__)


class ParsedDocument(BaseModel):
    """Result of document parsing."""

    text: str = Field(description="Extracted text content")
    metadata: dict[str, Any] = Field(default_factory=dict)
    pages: list[str] | None = Field(default=None, description="Text per page (if applicable)")
    page_count: int | None = Field(default=None)
    word_count: int = Field(default=0)
    char_count: int = Field(default=0)
    language: str | None = Field(default=None)
    error: str | None = Field(default=None)


class DocumentParser(ABC):
    """Abstract base class for document parsers."""

    @abstractmethod
    async def parse(self, content: bytes, filename: str) -> ParsedDocument:
        """Parse document content.

        Args:
            content: Document bytes.
            filename: Original filename (for format hints).

        Returns:
            Parsed document with extracted text.
        """
        ...

    @abstractmethod
    def supports(self, content_type: str) -> bool:
        """Check if parser supports the content type.

        Args:
            content_type: MIME type.

        Returns:
            True if parser can handle this type.
        """
        ...


class TextParser(DocumentParser):
    """Parser for plain text files."""

    SUPPORTED_TYPES = {
        "text/plain",
        "text/csv",
        "text/markdown",
        "text/html",
        "application/json",
        "application/xml",
        "text/xml",
    }

    async def parse(self, content: bytes, filename: str) -> ParsedDocument:
        """Parse text content."""
        try:
            # Try UTF-8 first, then fallback to latin-1
            try:
                text = content.decode("utf-8")
            except UnicodeDecodeError:
                text = content.decode("latin-1")

            # Clean up excessive whitespace
            text = "\n".join(line.rstrip() for line in text.splitlines())

            return ParsedDocument(
                text=text,
                word_count=len(text.split()),
                char_count=len(text),
                metadata={"encoding": "utf-8", "filename": filename},
            )
        except Exception as e:
            logger.warning("Text parsing failed", filename=filename, error=str(e))
            return ParsedDocument(text="", error=str(e))

    def supports(self, content_type: str) -> bool:
        return content_type in self.SUPPORTED_TYPES


class PDFParser(DocumentParser):
    """Parser for PDF documents."""

    async def parse(self, content: bytes, filename: str) -> ParsedDocument:
        """Parse PDF content."""
        try:
            from pypdf import PdfReader

            reader = PdfReader(io.BytesIO(content))

            pages = []
            full_text = []

            for page_num, page in enumerate(reader.pages):
                try:
                    text = page.extract_text() or ""
                    pages.append(text)
                    full_text.append(text)
                except Exception as e:
                    logger.warning(
                        "Failed to extract page",
                        page=page_num,
                        error=str(e),
                    )
                    pages.append("")

            text = "\n\n".join(full_text)

            # Extract metadata
            metadata: dict[str, Any] = {
                "filename": filename,
            }

            if reader.metadata:
                if reader.metadata.title:
                    metadata["title"] = reader.metadata.title
                if reader.metadata.author:
                    metadata["author"] = reader.metadata.author
                if reader.metadata.subject:
                    metadata["subject"] = reader.metadata.subject
                if reader.metadata.creation_date:
                    metadata["created"] = str(reader.metadata.creation_date)

            return ParsedDocument(
                text=text,
                pages=pages,
                page_count=len(reader.pages),
                word_count=len(text.split()),
                char_count=len(text),
                metadata=metadata,
            )
        except ImportError:
            return ParsedDocument(
                text="",
                error="pypdf not installed. Install with: pip install pypdf",
            )
        except Exception as e:
            logger.warning("PDF parsing failed", filename=filename, error=str(e))
            return ParsedDocument(text="", error=str(e))

    def supports(self, content_type: str) -> bool:
        return content_type == "application/pdf"


class DOCXParser(DocumentParser):
    """Parser for Microsoft Word documents."""

    SUPPORTED_TYPES = {
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        "application/msword",
    }

    async def parse(self, content: bytes, filename: str) -> ParsedDocument:
        """Parse DOCX content."""
        try:
            from docx import Document

            doc = Document(io.BytesIO(content))

            paragraphs = []
            for para in doc.paragraphs:
                text = para.text.strip()
                if text:
                    paragraphs.append(text)

            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(
                        cell.text.strip() for cell in row.cells if cell.text.strip()
                    )
                    if row_text:
                        paragraphs.append(row_text)

            text = "\n\n".join(paragraphs)

            # Extract metadata
            metadata: dict[str, Any] = {
                "filename": filename,
                "paragraph_count": len(doc.paragraphs),
                "table_count": len(doc.tables),
            }

            if doc.core_properties:
                if doc.core_properties.title:
                    metadata["title"] = doc.core_properties.title
                if doc.core_properties.author:
                    metadata["author"] = doc.core_properties.author
                if doc.core_properties.created:
                    metadata["created"] = str(doc.core_properties.created)

            return ParsedDocument(
                text=text,
                word_count=len(text.split()),
                char_count=len(text),
                metadata=metadata,
            )
        except ImportError:
            return ParsedDocument(
                text="",
                error="python-docx not installed. Install with: pip install python-docx",
            )
        except Exception as e:
            logger.warning("DOCX parsing failed", filename=filename, error=str(e))
            return ParsedDocument(text="", error=str(e))

    def supports(self, content_type: str) -> bool:
        return content_type in self.SUPPORTED_TYPES


class ImageParser(DocumentParser):
    """Parser for images (basic metadata, optional OCR)."""

    SUPPORTED_TYPES = {
        "image/jpeg",
        "image/png",
        "image/gif",
        "image/webp",
        "image/tiff",
        "image/bmp",
    }

    def __init__(self, enable_ocr: bool = False) -> None:
        """Initialize image parser.

        Args:
            enable_ocr: Enable OCR text extraction (requires tesseract).
        """
        self._enable_ocr = enable_ocr

    async def parse(self, content: bytes, filename: str) -> ParsedDocument:
        """Parse image, extracting metadata and optionally performing OCR."""
        try:
            from PIL import Image

            image = Image.open(io.BytesIO(content))

            metadata: dict[str, Any] = {
                "filename": filename,
                "format": image.format,
                "mode": image.mode,
                "width": image.width,
                "height": image.height,
            }

            # Extract EXIF if available
            if hasattr(image, "_getexif") and image._getexif():
                exif = image._getexif()
                if exif:
                    metadata["has_exif"] = True

            text = ""

            # Perform OCR if enabled
            if self._enable_ocr:
                try:
                    import pytesseract

                    text = pytesseract.image_to_string(image)
                    text = text.strip()
                    metadata["ocr_performed"] = True
                except ImportError:
                    metadata["ocr_error"] = "pytesseract not installed"
                except Exception as e:
                    metadata["ocr_error"] = str(e)

            return ParsedDocument(
                text=text,
                word_count=len(text.split()) if text else 0,
                char_count=len(text),
                metadata=metadata,
            )
        except ImportError:
            return ParsedDocument(
                text="",
                error="Pillow not installed. Install with: pip install Pillow",
            )
        except Exception as e:
            logger.warning("Image parsing failed", filename=filename, error=str(e))
            return ParsedDocument(text="", error=str(e))

    def supports(self, content_type: str) -> bool:
        return content_type in self.SUPPORTED_TYPES


class DocumentParserRegistry:
    """Registry for document parsers."""

    def __init__(self, enable_ocr: bool = False) -> None:
        """Initialize with default parsers.

        Args:
            enable_ocr: Enable OCR for image parsing.
        """
        self._parsers: list[DocumentParser] = [
            TextParser(),
            PDFParser(),
            DOCXParser(),
            ImageParser(enable_ocr=enable_ocr),
        ]

    def register(self, parser: DocumentParser) -> None:
        """Register a custom parser."""
        self._parsers.insert(0, parser)  # Custom parsers take priority

    def get_parser(self, content_type: str) -> DocumentParser | None:
        """Get parser for content type."""
        for parser in self._parsers:
            if parser.supports(content_type):
                return parser
        return None

    async def parse(
        self,
        content: bytes,
        content_type: str,
        filename: str,
    ) -> ParsedDocument:
        """Parse document using appropriate parser.

        Args:
            content: Document bytes.
            content_type: MIME type.
            filename: Original filename.

        Returns:
            Parsed document.

        Raises:
            ValueError: If no parser available for content type.
        """
        parser = self.get_parser(content_type)
        if not parser:
            # Try text parser as fallback for unknown types
            return await TextParser().parse(content, filename)

        return await parser.parse(content, filename)

    def supported_types(self) -> set[str]:
        """Get all supported content types."""
        types: set[str] = set()
        for parser in self._parsers:
            if isinstance(parser, TextParser):
                types.update(TextParser.SUPPORTED_TYPES)
            elif isinstance(parser, DOCXParser):
                types.update(DOCXParser.SUPPORTED_TYPES)
            elif isinstance(parser, ImageParser):
                types.update(ImageParser.SUPPORTED_TYPES)
            elif isinstance(parser, PDFParser):
                types.add("application/pdf")
        return types
